#!/usr/bin/env python3
"""Build multi-format reports from a compare_snapshots JSON report.

Creates HTML (default), CSV, XLSX (if openpyxl), DOCX (if python-docx), and PDF (wkhtmltopdf or weasyprint).
Output is placed under ~/Reports/<name>_<timestamp>/
"""
import argparse
import json
import os
import subprocess
import sys
from datetime import datetime


def ensure_dir(path):
    os.makedirs(path, exist_ok=True)


def call_html_generator(report_json, out_html):
    # call existing generator script
    script = os.path.join(os.path.dirname(__file__), 'generate_html_report.py')
    subprocess.check_call([sys.executable, script, report_json, out_html])


def write_summary_csv(report, out_dir):
    import csv
    out = os.path.join(out_dir, 'summary.csv')
    with open(out, 'w', newline='') as fh:
        w = csv.writer(fh)
        w.writerow(['item','added','removed','changed'])
        for k,v in report.get('summary', {}).items():
            if isinstance(v, dict):
                w.writerow([k, v.get('added',''), v.get('removed',''), v.get('changed','')])
            else:
                w.writerow([k, '', '', ''])
    print('Wrote', out)


def write_lists_csv(report, key, out_dir):
    import csv
    details = report.get('details', {})
    section = details.get(key, {})
    if not isinstance(section, dict):
        return
    out = os.path.join(out_dir, f'{key}.csv')
    with open(out, 'w', newline='') as fh:
        w = csv.writer(fh)
        w.writerow(['change','value'])
        for a in section.get('added', []):
            w.writerow(['added', a])
        for r in section.get('removed', []):
            w.writerow(['removed', r])
    print('Wrote', out)


def build_xlsx(report, out_dir):
    try:
        from openpyxl import Workbook
    except Exception:
        print('openpyxl not available; skipping xlsx')
        return
    wb = Workbook()
    ws = wb.active
    ws.title = 'Summary'
    ws.append(['Item','Added','Removed','Changed'])
    for k,v in report.get('summary', {}).items():
        if isinstance(v, dict):
            ws.append([k, v.get('added',''), v.get('removed',''), v.get('changed','')])
        else:
            ws.append([k,'','',''])
    # add sheets for packages/services
    details = report.get('details', {})
    for key in ['rpm_packages','enabled_services','running_services','users','groups']:
        sec = details.get(key, {})
        if isinstance(sec, dict):
            ws2 = wb.create_sheet(key)
            ws2.append(['change','value'])
            for a in sec.get('added', []):
                ws2.append(['added', a])
            for r in sec.get('removed', []):
                ws2.append(['removed', r])
    out = os.path.join(out_dir, 'report.xlsx')
    wb.save(out)
    print('Wrote', out)


def build_docx(report, out_dir):
    try:
        from docx import Document
    except Exception:
        print('python-docx not available; skipping docx')
        return
    doc = Document()
    doc.add_heading('PostInstall Healthcheck Report', level=1)
    doc.add_paragraph(f"Generated at: {report.get('generated_at','')}")
    doc.add_heading('Summary', level=2)
    for k,v in report.get('summary', {}).items():
        doc.add_paragraph(f"{k}: {v}")
    details = report.get('details', {})
    doc.add_heading('Details (sample)', level=2)
    for key in ['rpm_packages','enabled_services','running_services']:
        sec = details.get(key, {})
        if isinstance(sec, dict):
            doc.add_heading(key, level=3)
            if sec.get('added'):
                doc.add_paragraph('Added:')
                for a in sec.get('added'):
                    doc.add_paragraph(str(a), style='List Bullet')
            if sec.get('removed'):
                doc.add_paragraph('Removed:')
                for r in sec.get('removed'):
                    doc.add_paragraph(str(r), style='List Bullet')
    out = os.path.join(out_dir, 'report.docx')
    doc.save(out)
    print('Wrote', out)


def build_pdf_from_html(html_path, pdf_path):
    # Try wkhtmltopdf first
    if shutil_which('wkhtmltopdf'):
        try:
            subprocess.check_call(['wkhtmltopdf', html_path, pdf_path])
            print('Wrote', pdf_path)
            return
        except Exception:
            print('wkhtmltopdf failed; trying weasyprint')
    try:
        from weasyprint import HTML
        HTML(html_path).write_pdf(pdf_path)
        print('Wrote', pdf_path)
    except Exception:
        print('Could not generate PDF (no wkhtmltopdf or weasyprint available)')


def shutil_which(cmd):
    from shutil import which
    return which(cmd)


def main():
    p = argparse.ArgumentParser()
    p.add_argument('report_json')
    p.add_argument('--name', '-n', help='Healthcheck name (used for folder)', default=None)
    args = p.parse_args()

    report_json = args.report_json
    with open(report_json) as f:
        report = json.load(f)

    name = args.name or report.get('metadata', {}).get('inventory_hostname', 'healthcheck')
    ts = report.get('generated_at') or datetime.utcnow().isoformat()
    safe_ts = ts.replace(':','').replace('+','').replace('Z','')
    out_base = os.path.expanduser(os.path.join('~','Reports', f"{name}_{safe_ts}"))
    ensure_dir(out_base)

    # HTML
    html_out = os.path.join(out_base, 'report.html')
    try:
        call_html_generator(report_json, html_out)
    except Exception as e:
        print('HTML generation failed:', e)

    # CSV summary
    write_summary_csv(report, out_base)
    # lists
    for k in ['rpm_packages','enabled_services','running_services','users','groups']:
        write_lists_csv(report, k, out_base)

    # XLSX
    build_xlsx(report, out_base)

    # DOCX
    build_docx(report, out_base)

    # PDF
    try:
        build_pdf_from_html(html_out, os.path.join(out_base, 'report.pdf'))
    except Exception as e:
        print('PDF generation skipped or failed:', e)

    print('Reports available in', out_base)


if __name__ == '__main__':
    main()
